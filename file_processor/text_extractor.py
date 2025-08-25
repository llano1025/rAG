# text_extractor.py
from typing import Union, Optional
from pathlib import Path
import logging
import tempfile
import os
import subprocess
import platform
from io import BytesIO
from .type_detector import FileTypeDetector, UnsupportedFileType

# Optional imports with availability flags
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    logging.warning("PyMuPDF (fitz) not available, PDF processing will be limited")
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logging.warning("python-docx not available, DOCX processing will be limited")
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    logging.warning("BeautifulSoup4 not available, HTML processing will be limited")
    BeautifulSoup = None
    BS4_AVAILABLE = False

# Lazy import for OCR processor
def _get_ocr_processor():
    """Lazy import of OCRProcessor."""
    try:
        from .ocr_processor import OCRProcessor
        return OCRProcessor
    except ImportError as e:
        logging.warning(f"OCRProcessor not available: {e}")
        return None

class TextExtractor:
    """Extracts text content from various document formats."""
    
    def __init__(self):
        self.type_detector = FileTypeDetector()
        self.ocr_processor = None  # Will be initialized lazily
        self.logger = logging.getLogger(__name__)

    def extract(self, file_path: Union[str, Path], mime_type: Optional[str] = None) -> str:
        """Extract text from a document based on its type."""
        if mime_type is None:
            mime_type = self.type_detector.detect_type(file_path)

        extractors = {
            'application/pdf': self._extract_from_pdf,
            'application/msword': self._extract_from_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'text/plain': self._extract_from_text,
            'text/html': self._extract_from_html,
            'image/jpeg': self._extract_from_image,
            'image/png': self._extract_from_image,
            'image/tiff': self._extract_from_image,
            'image/gif': self._extract_from_image
        }

        extractor = extractors.get(mime_type)
        if not extractor:
            raise UnsupportedFileType(f"No text extractor for MIME type: {mime_type}")

        return extractor(file_path)

    async def extract_text(self, file_content: bytes, content_type: str, filename: str = None, 
                          ocr_method: str = None, ocr_language: str = None, vision_provider: str = None) -> str:
        """Extract text from document content bytes (async method expected by controllers)."""
        # Create temporary file from bytes content
        with tempfile.NamedTemporaryFile(delete=False, suffix=self._get_file_extension(content_type)) as temp_file:
            temp_file.write(file_content)
            temp_file.flush()
            temp_path = temp_file.name
        
        try:
            # For image files, use specified OCR settings
            if content_type and content_type.startswith('image/'):
                return self._extract_from_image(temp_path, ocr_method, ocr_language, vision_provider)
            else:
                # Use existing extract method for non-image files
                return self.extract(temp_path, content_type)
        finally:
            # Clean up temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    def _get_file_extension(self, content_type: str) -> str:
        """Get appropriate file extension for content type."""
        extensions = {
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'text/plain': '.txt',
            'text/html': '.html',
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/tiff': '.tiff',
            'image/gif': '.gif'
        }
        return extensions.get(content_type, '.tmp')

    def _extract_from_pdf(self, file_path: Union[str, Path]) -> str:
        """Extract text from PDF files."""
        if not PYMUPDF_AVAILABLE:
            self.logger.error("PyMuPDF not available for PDF extraction")
            return f"[Could not extract text from PDF file: {Path(file_path).name}. PyMuPDF (fitz) is required for PDF processing.]"
        
        try:
            text = []
            with fitz.open(str(file_path)) as doc:
                for page in doc:
                    text.append(page.get_text())
            return "\n".join(text)
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return f"[Error extracting text from PDF file: {str(e)}]"

    def _extract_from_docx(self, file_path: Union[str, Path]) -> str:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for DOCX extraction")
            return f"[Could not extract text from DOCX file: {Path(file_path).name}. python-docx is required for DOCX processing.]"
        
        try:
            doc = DocxDocument(str(file_path))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return f"[Error extracting text from DOCX file: {str(e)}]"

    def _extract_from_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from plain text files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _extract_from_html(self, file_path: Union[str, Path]) -> str:
        """Extract text from HTML files."""
        if not BS4_AVAILABLE:
            self.logger.warning("BeautifulSoup4 not available, using basic HTML text extraction")
            # Fallback: basic HTML tag removal
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                # Simple regex to remove HTML tags
                import re
                text = re.sub(r'<[^>]+>', '', html_content)
                # Clean up extra whitespace
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
            except Exception as e:
                self.logger.error(f"Failed to extract HTML text using fallback method: {e}")
                return f"[Error extracting text from HTML file: {str(e)}]"
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            self.logger.error(f"Failed to extract text from HTML {file_path}: {e}")
            return f"[Error extracting text from HTML file: {str(e)}]"

    def _extract_from_doc(self, file_path: Union[str, Path]) -> str:
        """Extract text from DOC files (legacy Word format)."""
        try:
            # Method 1: Try using python-docx2txt if available
            try:
                import docx2txt
                text = docx2txt.process(str(file_path))
                if text and text.strip():
                    return text
            except ImportError:
                logging.debug("docx2txt not available, trying alternative methods")
            
            # Method 2: Try using antiword (Unix/Linux systems)
            if platform.system() in ['Linux', 'Darwin']:  # Linux or macOS
                try:
                    result = subprocess.run(
                        ['antiword', str(file_path)],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        return result.stdout
                except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
                    logging.debug("antiword not available or failed")
            
            # Method 3: Try using catdoc (Unix/Linux systems)
            if platform.system() in ['Linux', 'Darwin']:
                try:
                    result = subprocess.run(
                        ['catdoc', str(file_path)],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        return result.stdout
                except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
                    logging.debug("catdoc not available or failed")
            
            # Method 4: Try using LibreOffice in headless mode (cross-platform)
            try:
                with tempfile.TemporaryDirectory() as temp_dir:
                    output_file = os.path.join(temp_dir, "output.txt")
                    
                    # Try common LibreOffice/OpenOffice paths
                    office_commands = [
                        'libreoffice',
                        'soffice',
                        '/usr/bin/libreoffice',
                        '/usr/local/bin/libreoffice',
                        'openoffice',
                        'ooffice'
                    ]
                    
                    for cmd in office_commands:
                        try:
                            result = subprocess.run([
                                cmd, '--headless', '--convert-to', 'txt:Text',
                                '--outdir', temp_dir, str(file_path)
                            ], capture_output=True, timeout=60, text=True)
                            
                            if result.returncode == 0:
                                # Look for output file
                                txt_file = Path(file_path).with_suffix('.txt')
                                txt_path = os.path.join(temp_dir, txt_file.name)
                                
                                if os.path.exists(txt_path):
                                    with open(txt_path, 'r', encoding='utf-8') as f:
                                        content = f.read()
                                    if content.strip():
                                        return content
                            break
                        except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
                            continue
            except Exception as e:
                logging.debug(f"LibreOffice conversion failed: {e}")
            
            # Method 5: Try using python-olefile for basic OLE structure parsing
            try:
                import olefile
                if olefile.isOleFile(str(file_path)):
                    with olefile.OleFileIO(str(file_path)) as ole:
                        # Try to find WordDocument stream
                        if ole._olestream_size is not None:
                            # This is a very basic attempt - OLE parsing is complex
                            # Better to use dedicated libraries
                            logging.debug("DOC file is valid OLE format but cannot extract text without proper parser")
            except ImportError:
                logging.debug("olefile not available")
            except Exception as e:
                logging.debug(f"OLE parsing failed: {e}")
            
            # Method 6: Fallback to binary text extraction (last resort)
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = content.decode(encoding, errors='ignore')
                        # Clean up binary data and extract readable text
                        cleaned_text = []
                        for line in text.split('\n'):
                            # Filter out lines with too much binary data
                            readable_chars = sum(1 for c in line if c.isprintable() or c.isspace())
                            if len(line) > 0 and readable_chars / len(line) > 0.7:
                                clean_line = ''.join(c for c in line if c.isprintable() or c.isspace())
                                if clean_line.strip():
                                    cleaned_text.append(clean_line.strip())
                        
                        result_text = '\n'.join(cleaned_text)
                        if result_text.strip():
                            logging.warning(f"Used fallback binary extraction for DOC file {file_path}")
                            return result_text
                    except Exception:
                        continue
            except Exception as e:
                logging.error(f"Binary fallback extraction failed: {e}")
            
            # If all methods fail, return informative message
            logging.warning(f"All DOC extraction methods failed for file {file_path}")
            return f"[Could not extract text from DOC file: {Path(file_path).name}. " \
                   f"Consider converting to DOCX format or install docx2txt, antiword, catdoc, or LibreOffice for better support.]"
                   
        except Exception as e:
            logging.error(f"Failed to extract text from DOC file {file_path}: {e}")
            return f"[Error extracting text from DOC file: {str(e)}]"

    def _extract_from_image(self, file_path: Union[str, Path], 
                            ocr_method: str = None, ocr_language: str = None, 
                            vision_provider: str = None) -> str:
        """Extract text from images using OCR with specified settings."""
        try:
            # Import OCR classes
            OCRProcessor = _get_ocr_processor()
            if not OCRProcessor:
                self.logger.error("OCRProcessor not available for image text extraction")
                return f"[Could not extract text from image file: {Path(file_path).name}. OCR dependencies are required for image processing.]"
            
            # Import enum classes for OCR configuration
            from .ocr_processor import OCRMethod, VisionProvider
            
            # Set defaults
            method = OCRMethod.TESSERACT
            language = ocr_language or 'eng'
            provider = VisionProvider.GEMINI
            
            # Parse OCR method
            if ocr_method:
                try:
                    method = OCRMethod(ocr_method)
                except ValueError:
                    self.logger.warning(f"Invalid OCR method '{ocr_method}', using Tesseract")
                    method = OCRMethod.TESSERACT
            
            # Parse vision provider
            if vision_provider and method == OCRMethod.VISION_LLM:
                try:
                    provider = VisionProvider(vision_provider)
                except ValueError:
                    self.logger.warning(f"Invalid vision provider '{vision_provider}', using Gemini")
                    provider = VisionProvider.GEMINI
            
            # Get vision LLM config if needed
            vision_llm_config = None
            if method == OCRMethod.VISION_LLM:
                vision_llm_config = self._get_vision_llm_config(provider)
                if not vision_llm_config:
                    self.logger.warning("Vision LLM config not available, falling back to Tesseract")
                    method = OCRMethod.TESSERACT
            
            # Create OCR processor with specific settings
            ocr_processor = OCRProcessor(
                method=method,
                language=language,
                vision_llm_config=vision_llm_config,
                vision_provider=provider,
                enable_fallback=True,
                timeout=60
            )
            
            return ocr_processor.process(file_path, vision_provider=provider if method == OCRMethod.VISION_LLM else None)
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from image {file_path}: {e}")
            return f"[Error extracting text from image file: {str(e)}]"

    def _get_vision_llm_config(self, provider: 'VisionProvider') -> dict:
        """Get vision LLM configuration for the specified provider."""
        try:
            from config import get_settings
            settings = get_settings()
            
            if provider.value == 'openai':
                if not getattr(settings, 'OPENAI_API_KEY', None):
                    return None
                return {
                    'api_key': settings.OPENAI_API_KEY,
                    'model': getattr(settings, 'OPENAI_MODEL', 'gpt-4o-2024-08-06'),
                    'endpoint': 'https://api.openai.com/v1/chat/completions',
                    'additional_params': {
                        'max_tokens': getattr(settings, 'OPENAI_MAX_TOKENS', 4096),
                        'temperature': getattr(settings, 'OPENAI_TEMPERATURE', 0.1)
                    }
                }
            elif provider.value == 'gemini':
                if not getattr(settings, 'GEMINI_API_KEY', None):
                    return None
                return {
                    'api_key': settings.GEMINI_API_KEY,
                    'model': getattr(settings, 'GEMINI_MODEL', 'gemini-2.5-flash'),
                    'additional_params': {
                        'max_tokens': getattr(settings, 'GEMINI_MAX_TOKENS', 8192),
                        'temperature': getattr(settings, 'GEMINI_TEMPERATURE', 0.1)
                    }
                }
            elif provider.value == 'claude':
                if not getattr(settings, 'ANTHROPIC_API_KEY', None):
                    return None
                return {
                    'api_key': settings.ANTHROPIC_API_KEY,
                    'model': getattr(settings, 'ANTHROPIC_MODEL', 'claude-3-5-sonnet-20241022'),
                    'endpoint': 'https://api.anthropic.com/v1/messages',
                    'additional_params': {
                        'max_tokens': getattr(settings, 'ANTHROPIC_MAX_TOKENS', 8192),
                        'temperature': getattr(settings, 'ANTHROPIC_TEMPERATURE', 0.1)
                    }
                }
        except Exception as e:
            self.logger.error(f"Failed to get vision LLM config: {e}")
        
        return None

    def extract_tables_from_pdf(self, file_path: Union[str, Path]) -> list:
        """Extract tables from PDF files."""
        try:
            import tabula
            # Extract tables using tabula-py
            tables = tabula.read_pdf(str(file_path), pages='all', multiple_tables=True)
            
            table_texts = []
            for i, table in enumerate(tables):
                if not table.empty:
                    # Convert table to text representation
                    table_text = f"Table {i+1}:\n"
                    table_text += table.to_string(index=False)
                    table_texts.append(table_text)
            
            return table_texts
            
        except ImportError:
            logging.debug("tabula-py not available for table extraction")
            return self._extract_tables_with_pymupdf(file_path)
        except Exception as e:
            logging.error(f"Table extraction failed with tabula: {e}")
            return self._extract_tables_with_pymupdf(file_path)

    def _extract_tables_with_pymupdf(self, file_path: Union[str, Path]) -> list:
        """Extract tables using PyMuPDF as fallback."""
        if not PYMUPDF_AVAILABLE:
            self.logger.warning("PyMuPDF not available for table extraction")
            return [f"[Could not extract tables from PDF file: {Path(file_path).name}. PyMuPDF is required for table processing.]"]
        
        try:
            table_texts = []
            with fitz.open(str(file_path)) as doc:
                for page_num, page in enumerate(doc):
                    # Find tables on the page
                    tables = page.find_tables()
                    
                    for table_num, table in enumerate(tables):
                        try:
                            # Extract table data
                            table_data = table.extract()
                            if table_data:
                                table_text = f"Page {page_num+1}, Table {table_num+1}:\n"
                                
                                # Convert table data to text
                                for row in table_data:
                                    if row:  # Skip empty rows
                                        # Clean and join cells
                                        clean_row = [str(cell).strip() if cell else "" for cell in row]
                                        table_text += " | ".join(clean_row) + "\n"
                                
                                table_texts.append(table_text)
                        except Exception as e:
                            logging.debug(f"Failed to extract table {table_num} from page {page_num}: {e}")
            
            return table_texts
            
        except Exception as e:
            logging.error(f"PyMuPDF table extraction failed: {e}")
            return []

    def extract_tables_from_docx(self, file_path: Union[str, Path]) -> list:
        """Extract tables from DOCX files."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for DOCX table extraction")
            return [f"[Could not extract tables from DOCX file: {Path(file_path).name}. python-docx is required for DOCX processing.]"]
        
        try:
            doc = DocxDocument(str(file_path))
            table_texts = []
            
            for table_num, table in enumerate(doc.tables):
                table_text = f"Table {table_num+1}:\n"
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if any(row_data):  # Skip empty rows
                        table_text += " | ".join(row_data) + "\n"
                
                table_texts.append(table_text)
            
            return table_texts
            
        except Exception as e:
            logging.error(f"DOCX table extraction failed: {e}")
            return []

    def extract_tables_from_html(self, file_path: Union[str, Path]) -> list:
        """Extract tables from HTML files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            tables = soup.find_all('table')
            table_texts = []
            
            for table_num, table in enumerate(tables):
                table_text = f"Table {table_num+1}:\n"
                
                rows = table.find_all('tr')
                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    row_data = [cell.get_text(strip=True) for cell in cells]
                    
                    if any(row_data):  # Skip empty rows
                        table_text += " | ".join(row_data) + "\n"
                
                table_texts.append(table_text)
            
            return table_texts
            
        except Exception as e:
            logging.error(f"HTML table extraction failed: {e}")
            return []

    def extract_with_tables(self, file_path: Union[str, Path], mime_type: Optional[str] = None) -> dict:
        """Extract both text and tables from a document."""
        if mime_type is None:
            mime_type = self.type_detector.detect_type(file_path)
        
        # Extract main text
        text = self.extract(file_path, mime_type)
        
        # Extract tables based on file type
        tables = []
        if mime_type == 'application/pdf':
            tables = self.extract_tables_from_pdf(file_path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            tables = self.extract_tables_from_docx(file_path)
        elif mime_type == 'text/html':
            tables = self.extract_tables_from_html(file_path)
        
        return {
            'text': text,
            'tables': tables,
            'table_count': len(tables)
        }

    def extract_structured_content(self, file_path: Union[str, Path], mime_type: Optional[str] = None) -> dict:
        """Extract structured content including text, tables, and metadata."""
        if mime_type is None:
            mime_type = self.type_detector.detect_type(file_path)
        
        result = {
            'text': '',
            'tables': [],
            'sections': [],
            'metadata': {
                'file_path': str(file_path),
                'mime_type': mime_type,
                'extraction_method': 'structured'
            }
        }
        
        try:
            # Extract with tables
            extraction_result = self.extract_with_tables(file_path, mime_type)
            result['text'] = extraction_result['text']
            result['tables'] = extraction_result['tables']
            
            # Extract sections for supported formats
            if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                result['sections'] = self._extract_docx_sections(file_path)
            elif mime_type == 'text/html':
                result['sections'] = self._extract_html_sections(file_path)
            
            result['metadata']['table_count'] = len(result['tables'])
            result['metadata']['section_count'] = len(result['sections'])
            result['metadata']['text_length'] = len(result['text'])
            
        except Exception as e:
            logging.error(f"Structured content extraction failed: {e}")
            result['metadata']['error'] = str(e)
        
        return result

    def _extract_docx_sections(self, file_path: Union[str, Path]) -> list:
        """Extract sections from DOCX based on heading styles."""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for DOCX section extraction")
            return []
        
        try:
            doc = DocxDocument(str(file_path))
            sections = []
            current_section = None
            
            for paragraph in doc.paragraphs:
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    if current_section:
                        sections.append(current_section)
                    
                    current_section = {
                        'title': paragraph.text.strip(),
                        'level': int(paragraph.style.name.replace('Heading ', '')) if paragraph.style.name.replace('Heading ', '').isdigit() else 1,
                        'content': ''
                    }
                elif current_section and paragraph.text.strip():
                    current_section['content'] += paragraph.text + '\n'
            
            # Add the last section
            if current_section:
                sections.append(current_section)
            
            return sections
            
        except Exception as e:
            logging.error(f"DOCX section extraction failed: {e}")
            return []

    def _extract_html_sections(self, file_path: Union[str, Path]) -> list:
        """Extract sections from HTML based on heading tags."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            
            sections = []
            current_section = None
            
            # Find all heading tags and content
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                if element.name.startswith('h'):
                    if current_section:
                        sections.append(current_section)
                    
                    current_section = {
                        'title': element.get_text(strip=True),
                        'level': int(element.name[1]),
                        'content': ''
                    }
                elif current_section and element.get_text(strip=True):
                    current_section['content'] += element.get_text(strip=True) + '\n'
            
            # Add the last section
            if current_section:
                sections.append(current_section)
            
            return sections
            
        except Exception as e:
            logging.error(f"HTML section extraction failed: {e}")
            return []
